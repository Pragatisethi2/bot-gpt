from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlmodel import Session, select
from app.database import get_session
from app.models import User, Conversation, Message, Document
from app.schemas import *
from app.services.llm_service import call_gemini_chat, call_gemini_rag, generate_summary
from app.services.rag_service import chunk_text, get_embeddings, retrieve_relevant_chunks
from datetime import datetime
import json
import PyPDF2
from docx import Document as DocxDocument
import io

router = APIRouter()
@router.post("/users", response_model=dict)
def create_user(name: str, email: str, db: Session = Depends(get_session)):
    # Check if user already exists
    existing_user = db.exec(select(User).where(User.email == email)).first()
    
    if existing_user:
        # Return existing user instead of creating new
        return {
            "user_id": existing_user.id,
            "name": existing_user.name,
            "email": existing_user.email,
            "message": "User already exists, logged in successfully"
        }
    
    # Create new user
    user = User(name=name, email=email)
    db.add(user)
    db.commit()
    db.refresh(user)
    
    return {
        "user_id": user.id,
        "name": user.name,
        "email": user.email,
        "message": "User created successfully"
    }

@router.get("/users/{user_id}", response_model=dict)
def get_user(user_id: int, db: Session = Depends(get_session)):
    user = db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    return {"user_id": user.id, "name": user.name, "email": user.email}
@router.post("/conversations", response_model=dict)
def create_conversation(request: CreateConversationRequest, db: Session = Depends(get_session)):
    conv = Conversation(
        user_id=request.user_id,
        title=request.first_message[:50],
        mode=request.mode
    )
    db.add(conv)
    db.commit()
    db.refresh(conv)
    
    user_msg = Message(conversation_id=conv.id, role="user", content=request.first_message)
    db.add(user_msg)
    db.commit()
    
    if request.mode == "chat":
        response_text = call_gemini_chat(None, [{"role": "user", "content": request.first_message}])
    else:
        response_text = "Please upload a document to start RAG conversation."
    
    model_msg = Message(conversation_id=conv.id, role="model", content=response_text)
    db.add(model_msg)
    db.commit()
    
    return {"conversation_id": conv.id, "response": response_text}

@router.get("/conversations", response_model=List[ConversationResponse])
def list_conversations(user_id: int, db: Session = Depends(get_session)):
    conversations = db.exec(
        select(Conversation).where(Conversation.user_id == user_id).order_by(Conversation.last_updated.desc())
    ).all()
    return conversations

@router.get("/conversations/{conv_id}", response_model=ConversationDetailResponse)
def get_conversation(conv_id: int, db: Session = Depends(get_session)):
    conv = db.get(Conversation, conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    messages = db.exec(
        select(Message).where(Message.conversation_id == conv_id).order_by(Message.timestamp)
    ).all()
    
    return {"conversation": conv, "messages": messages}

@router.post("/conversations/{conv_id}/messages", response_model=dict)
def add_message(conv_id: int, request: AddMessageRequest, db: Session = Depends(get_session)):
    conv = db.get(Conversation, conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    user_msg = Message(conversation_id=conv_id, role="user", content=request.content)
    db.add(user_msg)
    db.commit()
    
    all_messages = db.exec(
        select(Message).where(Message.conversation_id == conv_id).order_by(Message.timestamp)
    ).all()
    
    message_count = len(all_messages)
    
    if message_count % 15 == 0 and message_count > 0:
        start_idx = max(0, message_count - 15)
        messages_to_summarize = all_messages[start_idx:message_count]
        msg_dicts = [{"role": m.role, "content": m.content} for m in messages_to_summarize]
        new_summary = generate_summary(msg_dicts)
        
        if conv.summary:
            conv.summary = f"{conv.summary}\n\n{new_summary}"
        else:
            conv.summary = new_summary
        
        db.add(conv)
        db.commit()
    
    recent_messages = all_messages[-10:]
    recent_msg_dicts = [{"role": m.role, "content": m.content} for m in recent_messages]
    
    if conv.mode == "chat":
        response_text = call_gemini_chat(conv.summary, recent_msg_dicts)
    else:
        doc = db.exec(select(Document).where(Document.conversation_id == conv_id)).first()
        if not doc:
            return {"error": "No document uploaded for RAG mode"}
        
        chunks = json.loads(doc.chunks)
        embeddings = json.loads(doc.embeddings)
        context = retrieve_relevant_chunks(request.content, chunks, embeddings)
        response_text = call_gemini_rag(request.content, context, conv.summary)
    
    model_msg = Message(conversation_id=conv_id, role="model", content=response_text)
    db.add(model_msg)
    
    conv.last_updated = datetime.utcnow()
    db.add(conv)
    db.commit()
    
    return {"response": response_text}

@router.delete("/conversations/{conv_id}", response_model=dict)
def delete_conversation(conv_id: int, db: Session = Depends(get_session)):
    conv = db.get(Conversation, conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    db.exec(select(Message).where(Message.conversation_id == conv_id)).all()
    for msg in db.exec(select(Message).where(Message.conversation_id == conv_id)):
        db.delete(msg)
    
    db.exec(select(Document).where(Document.conversation_id == conv_id)).all()
    for doc in db.exec(select(Document).where(Document.conversation_id == conv_id)):
        db.delete(doc)
    
    db.delete(conv)
    db.commit()
    
    return {"status": "deleted"}

@router.post("/conversations/{conv_id}/documents", response_model=dict)
async def upload_document(
    conv_id: int, 
    file: UploadFile = File(...),
    title: str = Form(None),
    db: Session = Depends(get_session)
):
    conv = db.get(Conversation, conv_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    
    print(f"Uploading file: {file.filename}")
    file_content = await file.read()
    
    if file.filename.endswith('.pdf'):
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            print(f"Extracted {len(text)} characters from PDF")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")
    
    elif file.filename.endswith('.docx'):
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
            print(f"Extracted {len(text)} characters from DOCX")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading DOCX: {str(e)}")
    
    elif file.filename.endswith('.txt'):
        try:
            text = file_content.decode('utf-8')
            print(f"Extracted {len(text)} characters from TXT")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Error reading TXT: {str(e)}")
    
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type. Use PDF, DOCX, or TXT")
    
    if not text.strip():
        raise HTTPException(status_code=400, detail="Document is empty or could not extract text")
    
    print("Chunking text...")
    chunks = chunk_text(text)
    print(f"Created {len(chunks)} chunks")
    
    print("Generating embeddings...")
    embeddings = get_embeddings(chunks)
    print(f"Generated {len(embeddings)} embeddings")
    
    if not embeddings or len(embeddings) != len(chunks):
        raise HTTPException(status_code=500, detail="Failed to generate embeddings")
    
    doc_title = title if title else file.filename
    
    doc = Document(
        conversation_id=conv_id,
        title=doc_title,
        text=text[:1000],  # Store first 1000 chars only to save space
        chunks=json.dumps(chunks),
        embeddings=json.dumps(embeddings)
    )
    db.add(doc)
    db.commit()
    
    print("Document saved successfully!")
    
    return {
        "status": "uploaded", 
        "filename": file.filename,
        "title": doc_title,
        "chunks_count": len(chunks),
        "embeddings_count": len(embeddings),
        "text_length": len(text)
    }

@router.get("/conversations/{conv_id}/documents", response_model=list)
def get_conversation_documents(conv_id: int, db: Session = Depends(get_session)):
    """Get all documents for a conversation"""
    docs = db.exec(select(Document).where(Document.conversation_id == conv_id)).all()
    return [{"id": d.id, "title": d.title, "created_at": d.created_at} for d in docs]